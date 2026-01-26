"""
GraphRAG æ ¸å¿ƒå®ç° 

"""

import faiss
import numpy as np
from openai import OpenAI
from typing import List, Dict, Optional
import networkx as nx
from collections import defaultdict
import tiktoken
from community import community_louvain
import json
import pickle
from pathlib import Path
from datetime import datetime
import hashlib
import uuid
import logging

# æ–‡æ¡£è§£æåº“
from pypdf import PdfReader
import pdfplumber
from docx import Document as DocxDocument
import pandas as pd
import re
from config.settings import settings


logger = logging.getLogger(__name__)


class DocumentParser:
    """æ–‡æ¡£è§£æå™¨ï¼šæ”¯æŒ PDFã€DOCXã€TXTã€MDã€CSV ç­‰æ ¼å¼"""
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """è§£æ PDF æ–‡ä»¶"""
        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
                    
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables, 1):
                        if table:
                            table_text = f"\n[Table {table_num} on Page {page_num}]\n"
                            for row in table:
                                table_text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
                            text_parts.append(table_text)
                
                return "\n\n".join(text_parts)
        except Exception as e:
            logger.warning(f"pdfplumber è§£æå¤±è´¥ï¼Œä½¿ç”¨ pypdf: {e}")
            try:
                reader = PdfReader(file_path)
                text_parts = []
                for page_num, page in enumerate(reader.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_parts.append(f"[Page {page_num}]\n{text}")
                return "\n\n".join(text_parts)
            except Exception as e2:
                raise Exception(f"PDF è§£æå®Œå…¨å¤±è´¥: {e2}")
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """è§£æ DOCX æ–‡ä»¶"""
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table_num, table in enumerate(doc.tables, 1):
                table_text = f"\n[Table {table_num}]\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    table_text += row_text + "\n"
                text_parts.append(table_text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            raise Exception(f"DOCX è§£æå¤±è´¥: {e}")
    
    @staticmethod
    def parse_txt(file_path: str) -> str:
        """è§£æçº¯æ–‡æœ¬æ–‡ä»¶"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    @staticmethod
    def parse_markdown(file_path: str) -> str:
        """è§£æ Markdown æ–‡ä»¶"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    @staticmethod
    def parse_csv(file_path: str) -> str:
        """è§£æ CSV æ–‡ä»¶"""
        try:
            df = pd.read_csv(file_path)
            text = f"CSV Data ({len(df)} rows x {len(df.columns)} columns)\n\n"
            text += df.to_string(index=False)
            return text
        except Exception as e:
            raise Exception(f"CSV è§£æå¤±è´¥: {e}")
    
    @classmethod
    def parse_document(cls, file_path: str) -> str:
        """æ ¹æ®æ–‡ä»¶æ‰©å±•åè‡ªåŠ¨é€‰æ‹©è§£æå™¨"""
        path = Path(file_path)
        extension = path.suffix.lower()
        
        parsers = {
            '.pdf': cls.parse_pdf,
            '.docx': cls.parse_docx,
            '.doc': cls.parse_docx,
            '.txt': cls.parse_txt,
            '.md': cls.parse_markdown,
            '.markdown': cls.parse_markdown,
            '.csv': cls.parse_csv,
        }
        
        parser = parsers.get(extension)
        if parser is None:
            raise ValueError(f"ä¸æ”¯æŒçš„æ–‡ä»¶æ ¼å¼: {extension}")
        
        return parser(str(path))


class TextChunker:
    """æ™ºèƒ½æ–‡æœ¬åˆ†å—å™¨"""
    
    def __init__(self, chunk_size: int = 600, overlap: int = 100):
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.encoding = tiktoken.get_encoding("cl100k_base")
    
    def chunk_by_sentences(self, text: str) -> List[str]:
        """æŒ‰å¥å­åˆ†å—ï¼ˆä¿æŒå¥å­å®Œæ•´æ€§ï¼‰"""
        sentences = re.split(r'[.!?ã€‚ï¼ï¼Ÿ]\s+', text)
        
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_words = len(sentence.split())
            
            if current_length + sentence_words > self.chunk_size and current_chunk:
                chunks.append(' '.join(current_chunk))
                
                overlap_sentences = []
                overlap_length = 0
                for s in reversed(current_chunk):
                    s_length = len(s.split())
                    if overlap_length + s_length <= self.overlap:
                        overlap_sentences.insert(0, s)
                        overlap_length += s_length
                    else:
                        break
                
                current_chunk = overlap_sentences
                current_length = overlap_length
            
            current_chunk.append(sentence)
            current_length += sentence_words
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks
    
    def chunk_text(self, text: str) -> List[str]:
        """åˆ†å—æ–‡æœ¬"""
        return self.chunk_by_sentences(text)


class GraphRAGPipeline:
    """
    GraphRAG Pipeline
    
    1. æ·»åŠ æ–‡æ¡£åè‡ªåŠ¨ä¿å­˜
    2. åˆå§‹åŒ–æ—¶è‡ªåŠ¨åŠ è½½å·²æœ‰æ•°æ®
    3. text_chunks ä¸ºç©ºçš„é—®é¢˜
    """

    def __init__(self, llm_api_key: str, embedding_api_key: str, llm_url: str, 
                 embedding_url: str, embedding_name: str, embedding_dim: int,
                 llm_name: str, storage_dir: str = "./graphrag_storage"):
        
        self.llm_client = OpenAI(base_url=llm_url, api_key=llm_api_key)
        self.embedding_client = OpenAI(base_url=embedding_url, api_key=embedding_api_key)

        self.embedding_name = embedding_name
        self.llm_name = llm_name
        self.dimension = embedding_dim
        
        # å­˜å‚¨ç›®å½•
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        
        # æ–‡æ¡£ç®¡ç†
        self.document_parser = DocumentParser()
        self.text_chunker = TextChunker()
        self.documents = {}
        
        # UUID æ˜ å°„
        self.uuid_to_docid = {}
        self.docid_to_uuid = {}
        
        # å›¾è°±æ•°æ®
        self.text_chunks = []
        self.chunk_to_doc = {}
        self.entities = {}
        self.relationships = []
        self.claims = []
        
        # çŸ¥è¯†å›¾è°±
        self.graph = nx.Graph()
        
        # ç¤¾åŒºç»“æ„
        self.communities = {}
        self.community_summaries = {}
        
        # FAISS ç´¢å¼•
        self.community_summary_index = None
        self.community_embeddings = []
        
        self.encoding = tiktoken.get_encoding("cl100k_base")
        
        self._auto_load_if_exists()
    
    def _auto_load_if_exists(self):
        """åˆå§‹åŒ–æ—¶è‡ªåŠ¨åŠ è½½å·²æœ‰æ•°æ®"""
        try:
            self.load("default")
            logger.info(f"âœ… è‡ªåŠ¨åŠ è½½çŸ¥è¯†åº“æˆåŠŸ: {len(self.documents)} ä¸ªæ–‡æ¡£, {len(self.text_chunks)} ä¸ªchunks")
        except FileNotFoundError:
            logger.info("ğŸ“ æœªæ‰¾åˆ°å·²æœ‰çŸ¥è¯†åº“ï¼Œå°†åˆ›å»ºæ–°çš„")
        except Exception as e:
            logger.warning(f"âš ï¸ åŠ è½½çŸ¥è¯†åº“å¤±è´¥: {e}")
    
    # ==================== æ–‡æ¡£ç®¡ç† ====================
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """è®¡ç®—æ–‡ä»¶å“ˆå¸Œ"""
        with open(file_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()
    
    def add_document(self, file_path: str, metadata: Optional[Dict] = None, 
                    doc_uuid: Optional[str] = None) -> str:
        """
        æ·»åŠ æ–‡æ¡£
        
        Returns:
            æ–‡æ¡£çš„ UUID
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"æ–‡ä»¶ä¸å­˜åœ¨: {file_path}")
        
        if doc_uuid is None:
            doc_uuid = str(uuid.uuid4())
        
        file_hash = self._calculate_file_hash(str(file_path))
        
        # æ£€æŸ¥æ–‡æ¡£æ˜¯å¦å·²å­˜åœ¨
        if file_hash in self.documents:
            logger.info(f"âš ï¸ æ–‡æ¡£å·²å­˜åœ¨: {file_path.name}")
            return self.docid_to_uuid.get(file_hash, doc_uuid)
        
        logger.info(f"ğŸ“„ æ·»åŠ æ–‡æ¡£: {file_path.name} (UUID: {doc_uuid})")
        
        # è§£ææ–‡æ¡£
        text = self.document_parser.parse_document(str(file_path))
        logger.info(f"  ğŸ“ æ–‡æ¡£è§£æå®Œæˆï¼Œæ–‡æœ¬é•¿åº¦: {len(text)} å­—ç¬¦")
        
        # åˆ†å—
        chunks = self.text_chunker.chunk_text(text)
        logger.info(f"  âœ‚ï¸ åˆ†å—å®Œæˆ: {len(chunks)} ä¸ªchunks")
        
        # è®°å½•æ–‡æ¡£ä¿¡æ¯
        doc_info = {
            'uuid': doc_uuid,
            'path': str(file_path),
            'name': file_path.name,
            'hash': file_hash,
            'chunks': chunks,
            'chunk_ids': [],
            'metadata': metadata or {},
            'added_at': datetime.now().isoformat()
        }
        
        self.documents[file_hash] = doc_info
        self.uuid_to_docid[doc_uuid] = file_hash
        self.docid_to_uuid[file_hash] = doc_uuid
        
        # â˜…â˜…â˜…æå–å›¾å…ƒç´ å¹¶æ·»åŠ åˆ° text_chunks â˜…â˜…â˜…
        chunk_start_id = len(self.text_chunks)
        logger.info(f"  ğŸ” å¼€å§‹æå–å›¾å…ƒç´  (èµ·å§‹ID: {chunk_start_id})...")
        
        for chunk_id, chunk in enumerate(chunks):
            global_chunk_id = chunk_start_id + chunk_id
            
            # æå–å›¾å…ƒç´ 
            logger.info(f"    å¤„ç† chunk {chunk_id + 1}/{len(chunks)}...")
            elements = self.extract_graph_elements(chunk, global_chunk_id)
            
            # â˜… å…³é”®ï¼šæ·»åŠ åˆ° text_chunks
            self.text_chunks.append(elements)
            self.chunk_to_doc[global_chunk_id] = file_hash
            doc_info['chunk_ids'].append(global_chunk_id)
            
            logger.info(f"      æå–: {len(elements.get('entities', []))} å®ä½“, "
                       f"{len(elements.get('relationships', []))} å…³ç³»")
        
        logger.info(f"  âœ… å®Œæˆ: æå–äº† {len(chunks)} ä¸ªæ–‡æœ¬å—")
        logger.info(f"  ğŸ“Š å½“å‰æ€»è®¡: {len(self.text_chunks)} ä¸ªchunks")
        
        # â˜…â˜…â˜… æ·»åŠ æ–‡æ¡£åè‡ªåŠ¨ä¿å­˜ â˜…â˜…â˜…
        try:
            self.save("default")
            logger.info(f"  ğŸ’¾ çŸ¥è¯†åº“å·²è‡ªåŠ¨ä¿å­˜")
        except Exception as e:
            logger.error(f"  âŒ è‡ªåŠ¨ä¿å­˜å¤±è´¥: {e}")
        
        return doc_uuid
    
    def remove_document(self, doc_id: str):
        """åˆ é™¤æ–‡æ¡£ï¼ˆæ”¯æŒ UUID æˆ–å†…éƒ¨ IDï¼‰"""
        if doc_id in self.uuid_to_docid:
            internal_doc_id = self.uuid_to_docid[doc_id]
            doc_uuid = doc_id
        elif doc_id in self.documents:
            internal_doc_id = doc_id
            doc_uuid = self.docid_to_uuid.get(doc_id)
        else:
            raise ValueError(f"æ–‡æ¡£ä¸å­˜åœ¨: {doc_id}")
        
        logger.info(f"ğŸ—‘ï¸ åˆ é™¤æ–‡æ¡£: {self.documents[internal_doc_id]['name']}")
        
        # æ ‡è®°åˆ é™¤çš„ chunks
        chunk_ids = set(self.documents[internal_doc_id]['chunk_ids'])
        for chunk_id in chunk_ids:
            if chunk_id < len(self.text_chunks):
                self.text_chunks[chunk_id] = {'entities': [], 'relationships': [], 'claims': []}
            self.chunk_to_doc.pop(chunk_id, None)
        
        # åˆ é™¤æ˜ å°„
        if doc_uuid:
            self.uuid_to_docid.pop(doc_uuid, None)
            self.docid_to_uuid.pop(internal_doc_id, None)
        
        del self.documents[internal_doc_id]
        
        # è‡ªåŠ¨ä¿å­˜
        try:
            self.save("default")
            logger.info("  ğŸ’¾ åˆ é™¤åå·²è‡ªåŠ¨ä¿å­˜")
        except Exception as e:
            logger.error(f"  âŒ è‡ªåŠ¨ä¿å­˜å¤±è´¥: {e}")
        
        logger.info("  âœ… æ–‡æ¡£å·²åˆ é™¤")
    
    def list_documents(self) -> List[Dict]:
        """åˆ—å‡ºæ‰€æœ‰æ–‡æ¡£"""
        return [
            {
                'uuid': info['uuid'],
                'id': doc_id,
                'name': info['name'],
                'path': info['path'],
                'chunks': len(info['chunks']),
                'added_at': info['added_at'],
                'metadata': info['metadata']
            }
            for doc_id, info in self.documents.items()
        ]
    
    # ==================== å›¾å…ƒç´ æå– ====================
    
    def extract_graph_elements(self, text: str, chunk_id: int) -> Dict:
        """ä»æ–‡æœ¬æå–å›¾å…ƒç´ """
        
        prompt = f"""ä»ä»¥ä¸‹æ–‡æœ¬ä¸­æå–ç»“æ„åŒ–ä¿¡æ¯ï¼Œè¿”å›JSONæ ¼å¼ã€‚

æ–‡æœ¬:
{text}

æå–å†…å®¹:
1. entities: [{{"name": "å®ä½“å", "type": "ç±»å‹", "description": "æè¿°"}}]
2. relationships: [{{"source": "æºå®ä½“", "target": "ç›®æ ‡å®ä½“", "description": "å…³ç³»", "strength": 1-10}}]
3. claims: [{{"subject": "ä¸»ä½“", "object": "å®¢ä½“", "type": "FACT/OPINION", "description": "æè¿°", "date": "æ—¶é—´"}}]

åªè¿”å›JSONï¼Œä¸è¦å…¶ä»–å†…å®¹ã€‚
"""

        try:
            response = self.llm_client.chat.completions.create(
                model=self.llm_name,
                messages=[
                    {"role": "system", "content": "ä½ æ˜¯çŸ¥è¯†å›¾è°±ä¸“å®¶ã€‚"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0,
                response_format={"type": "json_object"}
            )
            
            result = json.loads(response.choices[0].message.content)
            logger.debug(f"Chunk {chunk_id} æå–ç»“æœ: {len(result.get('entities', []))} å®ä½“")
            return result
            
        except Exception as e:
            logger.error(f"æå–å¤±è´¥ (chunk {chunk_id}): {e}")
            return {"entities": [], "relationships": [], "claims": []}
    
    def summarize_entity(self, entity_name: str, descriptions: List[str]) -> str:
        """åˆå¹¶å®ä½“æè¿°"""
        if len(descriptions) == 1:
            return descriptions[0]
        
        combined = "\n".join([f"- {desc}" for desc in descriptions])
        
        prompt = f"""æ•´åˆä»¥ä¸‹å…³äº"{entity_name}"çš„æè¿°ä¸ºä¸€ä¸ªæ‘˜è¦ï¼ˆ150-200è¯ï¼‰ï¼š

{combined}

åªè¿”å›æ‘˜è¦ã€‚"""

        response = self.llm_client.chat.completions.create(
            model=self.llm_name,
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=300
        )
        
        return response.choices[0].message.content.strip()
    
    # ==================== å›¾è°±æ„å»º ====================
    
    def merge_entities_and_relationships(self):
        """åˆå¹¶å®ä½“å’Œå…³ç³»"""
        logger.info(f"ğŸ“Š å¼€å§‹åˆå¹¶å®ä½“å’Œå…³ç³» (text_chunksæ•°é‡: {len(self.text_chunks)})...")
        
        self.entities = {}
        self.relationships = []
        
        entity_descriptions = defaultdict(list)
        entity_types = {}
        entity_sources = defaultdict(set)
        
        # â˜…â˜…â˜…  text_chunks æ˜¯å¦ä¸ºç©º â˜…â˜…â˜…
        if not self.text_chunks:
            logger.warning("âš ï¸ text_chunks ä¸ºç©ºï¼è¯·å…ˆæ·»åŠ æ–‡æ¡£ã€‚")
            return
        
        for chunk_id, chunk_data in enumerate(self.text_chunks):
            entities = chunk_data.get('entities', [])
            logger.debug(f"  Chunk {chunk_id}: {len(entities)} å®ä½“")
            
            for entity in entities:
                name = entity['name']
                entity_descriptions[name].append(entity['description'])
                entity_types[name] = entity['type']
                entity_sources[name].add(chunk_id)
        
        logger.info(f"  å‘ç° {len(entity_descriptions)} ä¸ªå”¯ä¸€å®ä½“")
        logger.info("  ç”Ÿæˆå®ä½“æ‘˜è¦...")
        
        for entity_name, descriptions in entity_descriptions.items():
            summary = self.summarize_entity(entity_name, descriptions)
            self.entities[entity_name] = {
                'description': summary,
                'type': entity_types[entity_name],
                'source_ids': list(entity_sources[entity_name])
            }
        
        # åˆå¹¶å…³ç³»
        relationship_map = defaultdict(lambda: {'descriptions': [], 'strengths': [], 'sources': set()})
        
        for chunk_id, chunk_data in enumerate(self.text_chunks):
            for rel in chunk_data.get('relationships', []):
                key = (rel['source'], rel['target'])
                relationship_map[key]['descriptions'].append(rel['description'])
                relationship_map[key]['strengths'].append(rel.get('strength', 5))
                relationship_map[key]['sources'].add(chunk_id)
        
        for (source, target), data in relationship_map.items():
            if source in self.entities and target in self.entities:
                self.relationships.append({
                    'source': source,
                    'target': target,
                    'description': '; '.join(data['descriptions']),
                    'weight': float(np.mean(data['strengths'])),
                    'source_ids': list(data['sources'])
                })
        
        logger.info(f"  âœ… å®Œæˆ: {len(self.entities)} å®ä½“, {len(self.relationships)} å…³ç³»")
    
    def build_graph(self):
        """æ„å»ºçŸ¥è¯†å›¾è°±"""
        logger.info("ğŸ•¸ï¸ æ„å»ºçŸ¥è¯†å›¾è°±...")
        self.graph = nx.Graph()
        
        for entity_name, entity_data in self.entities.items():
            self.graph.add_node(
                entity_name,
                type=entity_data['type'],
                description=entity_data['description']
            )
        
        for rel in self.relationships:
            self.graph.add_edge(
                rel['source'],
                rel['target'],
                weight=rel['weight'],
                description=rel['description']
            )
        
        logger.info(f"  âœ… å›¾è°±: {self.graph.number_of_nodes()} èŠ‚ç‚¹, {self.graph.number_of_edges()} è¾¹")
    
    def detect_hierarchical_communities(self, max_level: int = 3):
        """å±‚æ¬¡åŒ–ç¤¾åŒºæ£€æµ‹"""
        logger.info("ğŸ‘¥ ç¤¾åŒºæ£€æµ‹...")
        
        self.communities = {}
        current_graph = self.graph.copy()
        
        for level in range(max_level):
            partition = community_louvain.best_partition(
                current_graph,
                weight='weight',
                resolution=1.0
            )
            
            communities_at_level = defaultdict(list)
            for node, comm_id in partition.items():
                communities_at_level[comm_id].append(node)
            
            self.communities[level] = dict(communities_at_level)
            logger.info(f"  Level {level}: {len(communities_at_level)} ä¸ªç¤¾åŒº")
            
            if len(communities_at_level) <= 1:
                break
            
            # æ„å»ºä¸‹ä¸€å±‚
            next_graph = nx.Graph()
            for comm_id in communities_at_level.keys():
                next_graph.add_node(f"comm_{level}_{comm_id}")
            
            for u, v, data in current_graph.edges(data=True):
                comm_u = partition[u]
                comm_v = partition[v]
                if comm_u != comm_v:
                    edge_key = (f"comm_{level}_{comm_u}", f"comm_{level}_{comm_v}")
                    if next_graph.has_edge(*edge_key):
                        next_graph[edge_key[0]][edge_key[1]]['weight'] += data.get('weight', 1)
                    else:
                        next_graph.add_edge(*edge_key, weight=data.get('weight', 1))
            
            current_graph = next_graph
    
    def generate_community_summary(self, level: int, community_id: int) -> str:
        """ç”Ÿæˆç¤¾åŒºæ‘˜è¦"""
        nodes = self.communities[level][community_id]
        
        entities_info = []
        for node in nodes[:20]:
            if node in self.entities:
                entities_info.append(
                    f"- {node} ({self.entities[node]['type']}): "
                    f"{self.entities[node]['description'][:200]}"
                )
        
        relationships_info = []
        for rel in self.relationships:
            if rel['source'] in nodes and rel['target'] in nodes:
                relationships_info.append(
                    f"- {rel['source']} â†’ {rel['target']}: {rel['description'][:150]}"
                )
        
        prompt = f"""ç”Ÿæˆç¤¾åŒºæ‘˜è¦ï¼ˆ300-400è¯ï¼‰ï¼š

å®ä½“:
{chr(10).join(entities_info)}

å…³ç³»:
{chr(10).join(relationships_info[:15])}

åŒ…æ‹¬ï¼šä¸»é¢˜ã€å…³é”®å®ä½“ã€å…³é”®å‘ç°ã€è¿æ¥æ€§ã€‚åªè¿”å›æ‘˜è¦ã€‚"""

        response = self.llm_client.chat.completions.create(
            model=self.llm_name,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=500
        )
        
        return response.choices[0].message.content.strip()
    
    def generate_all_community_summaries(self):
        """ç”Ÿæˆæ‰€æœ‰ç¤¾åŒºæ‘˜è¦"""
        logger.info("ğŸ“ ç”Ÿæˆç¤¾åŒºæ‘˜è¦...")
        self.community_summaries = {}
        
        for level, communities in self.communities.items():
            for comm_id in communities.keys():
                summary = self.generate_community_summary(level, comm_id)
                self.community_summaries[(level, comm_id)] = summary
        
        logger.info(f"  âœ… ç”Ÿæˆäº† {len(self.community_summaries)} ä¸ªç¤¾åŒºæ‘˜è¦")
    
    def build_community_summary_index(self):
        """æ„å»ºå‘é‡ç´¢å¼•"""
        logger.info("ğŸ” æ„å»ºå‘é‡ç´¢å¼•...")
        
        summaries = []
        summary_metadata = []
        
        for (level, comm_id), summary in self.community_summaries.items():
            summaries.append(summary)
            summary_metadata.append({
                'level': level,
                'community_id': comm_id,
                'summary': summary
            })
        
        if not summaries:
            logger.warning("âš ï¸ æ²¡æœ‰ç¤¾åŒºæ‘˜è¦å¯ç´¢å¼•")
            return
        
        # ç”Ÿæˆ embeddings
        embeddings = []
        batch_size = 100
        
        logger.info(f"  ç”Ÿæˆ {len(summaries)} ä¸ªæ‘˜è¦çš„å‘é‡...")
        for i in range(0, len(summaries), batch_size):
            batch = summaries[i:i + batch_size]
            response = self.embedding_client.embeddings.create(
                model=self.embedding_name,
                input=batch
            )
            batch_embeddings = [np.array(item.embedding, dtype='float32') 
                              for item in response.data]
            embeddings.extend(batch_embeddings)
        
        self.community_embeddings = summary_metadata
        
        # æ„å»º FAISS
        embeddings_array = np.array(embeddings, dtype='float32')
        self.community_summary_index = faiss.IndexFlatIP(self.dimension)
        faiss.normalize_L2(embeddings_array)
        self.community_summary_index.add(embeddings_array)
        
        logger.info(f"  âœ… ç´¢å¼•å®Œæˆ: {len(embeddings)} ä¸ªç¤¾åŒº")
    
    # ==================== ç´¢å¼•æ„å»º ====================
    
    def rebuild_index(self):
        """é‡å»ºç´¢å¼•"""
        logger.info("=" * 60)
        logger.info("ğŸ”„ é‡å»º GraphRAG ç´¢å¼•")
        logger.info("=" * 60)
        
        # â˜…â˜…â˜… æ£€æŸ¥æ˜¯å¦æœ‰æ•°æ® â˜…â˜…â˜…
        if not self.text_chunks:
            logger.error("âŒ text_chunks ä¸ºç©ºï¼è¯·å…ˆæ·»åŠ æ–‡æ¡£ã€‚")
            raise RuntimeError("æ²¡æœ‰æ–‡æ¡£å¯ä»¥ç´¢å¼•ï¼Œè¯·å…ˆä¸Šä¼ æ–‡æ¡£")
        
        logger.info(f"ğŸ“Š æ•°æ®ç»Ÿè®¡: {len(self.text_chunks)} chunks, {len(self.documents)} æ–‡æ¡£")
        
        logger.info("[1/5] åˆå¹¶å®ä½“å’Œå…³ç³»...")
        self.merge_entities_and_relationships()
        logger.info(f"  å®Œæˆ: {len(self.entities)} å®ä½“, {len(self.relationships)} å…³ç³»")
        
        if not self.entities:
            logger.error("âŒ æ²¡æœ‰æå–åˆ°å®ä½“ï¼è¯·æ£€æŸ¥æ–‡æ¡£å†…å®¹æˆ– LLM é…ç½®")
            raise RuntimeError("æœªèƒ½æå–å®ä½“ï¼Œç´¢å¼•æ„å»ºå¤±è´¥")
        
        logger.info("[2/5] æ„å»ºçŸ¥è¯†å›¾è°±...")
        self.build_graph()
        logger.info(f"  å®Œæˆ: {self.graph.number_of_nodes()} èŠ‚ç‚¹, {self.graph.number_of_edges()} è¾¹")
        
        logger.info("[3/5] ç¤¾åŒºæ£€æµ‹...")
        self.detect_hierarchical_communities()
        
        logger.info("[4/5] ç”Ÿæˆç¤¾åŒºæ‘˜è¦...")
        self.generate_all_community_summaries()
        
        logger.info("[5/5] æ„å»ºå‘é‡ç´¢å¼•...")
        self.build_community_summary_index()
        
        # â˜…â˜…â˜…é‡å»ºç´¢å¼•åè‡ªåŠ¨ä¿å­˜ â˜…â˜…â˜…
        try:
            self.save("default")
            logger.info("ğŸ’¾ ç´¢å¼•é‡å»ºåå·²è‡ªåŠ¨ä¿å­˜")
        except Exception as e:
            logger.error(f"âŒ è‡ªåŠ¨ä¿å­˜å¤±è´¥: {e}")
        
        logger.info("=" * 60)
        logger.info("âœ… ç´¢å¼•é‡å»ºå®Œæˆ!")
        logger.info("=" * 60)
    
    # ==================== æŸ¥è¯¢ ====================
    
    def global_query(self, question: str, top_k_communities: int = 10, return_sample=False) -> str:
        """æŸ¥è¯¢çŸ¥è¯†åº“"""
        if self.community_summary_index is None:
            raise RuntimeError("ç´¢å¼•æœªæ„å»ºï¼Œè¯·å…ˆä¸Šä¼ æ–‡æ¡£å¹¶é‡å»ºç´¢å¼•")
        
        # æ£€ç´¢ç¤¾åŒº
        query_embedding = self._get_embedding(question)
        query_embedding = np.array([query_embedding], dtype='float32')
        faiss.normalize_L2(query_embedding)
        
        scores, indices = self.community_summary_index.search(
            query_embedding, 
            min(top_k_communities, len(self.community_embeddings))
        )
        
        # ç®€å•æ¨¡å¼ï¼šç›´æ¥è¿”å›ç¤¾åŒºæ‘˜è¦
        if getattr(settings, 'SIMPLE_RAG', False) or return_sample:
            search_results = []
            threshold = getattr(settings, 'T_SCORE', 0.5)
            
            for idx, score in zip(indices[0], scores[0]):
                if score >= threshold:
                    search_results.append(self.community_embeddings[idx]['summary'])
            
            if not search_results:
                return "æŠ±æ­‰ï¼Œæœªæ‰¾åˆ°ç›¸å…³ä¿¡æ¯ã€‚"
            
            end_str = ""
            for i, res in enumerate(search_results):
                end_str += f"ç¤¾åŒºæ‘˜è¦ {i+1}\n{res}\n\n"
            return end_str
        
        # Map-Reduce æ¨¡å¼
        community_answers = []
        
        for idx, score in zip(indices[0], scores[0]):
            if idx >= len(self.community_embeddings):
                continue
            
            comm_data = self.community_embeddings[idx]
            answer = self._ask_community(question, comm_data['summary'])
            
            if answer and len(answer.strip()) > 10:
                community_answers.append({
                    'level': comm_data['level'],
                    'community_id': comm_data['community_id'],
                    'content': answer,
                    'score': float(score)
                })
        
        return self._reduce_answers(question, community_answers)
    
    def _ask_community(self, question: str, community_summary: str) -> str:
        """è¯¢é—®å•ä¸ªç¤¾åŒº"""
        prompt = f"""åŸºäºç¤¾åŒºä¿¡æ¯å›ç­”é—®é¢˜ï¼ˆ2-3å¥è¯ï¼‰ã€‚å¦‚æœæ— å…³ï¼Œå›ç­”"æ— ç›¸å…³ä¿¡æ¯"ã€‚

ç¤¾åŒºä¿¡æ¯:
{community_summary}

é—®é¢˜: {question}

åªè¿”å›ç­”æ¡ˆã€‚"""
        
        try:
            response = self.llm_client.chat.completions.create(
                model=self.llm_name,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.5,
                max_tokens=150
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"ç¤¾åŒºæŸ¥è¯¢å¤±è´¥: {e}")
            return ""
    
    def _reduce_answers(self, question: str, community_answers: List[Dict]) -> str:
        """ç»¼åˆç­”æ¡ˆ"""
        if not community_answers:
            return "æŠ±æ­‰ï¼Œåœ¨çŸ¥è¯†å›¾è°±ä¸­æ²¡æœ‰æ‰¾åˆ°ç›¸å…³ä¿¡æ¯ã€‚"
        
        community_answers.sort(key=lambda x: x['score'], reverse=True)
        
        answers_text = []
        for i, ans_data in enumerate(community_answers[:10], 1):
            if ans_data['content'].lower() != "æ— ç›¸å…³ä¿¡æ¯":
                answers_text.append(f"{i}. {ans_data['content']}")
        
        if not answers_text:
            return "æŠ±æ­‰ï¼Œæ‰¾åˆ°çš„ä¿¡æ¯ä¸é—®é¢˜ä¸å¤ªç›¸å…³ã€‚"
        
        combined = "\n".join(answers_text)
        
        prompt = f"""ç»¼åˆä»¥ä¸‹ç­”æ¡ˆä¸ºä¸€ä¸ªè¿è´¯çš„æœ€ç»ˆç­”æ¡ˆï¼ˆ200-400è¯ï¼‰ï¼š

é—®é¢˜: {question}

å„ç¤¾åŒºç­”æ¡ˆ:
{combined}

è¦æ±‚: æ•´åˆä¿¡æ¯ã€æ¶ˆé™¤å†—ä½™ã€ä¿æŒæ¸…æ™°ã€å‘ˆç°ä¸åŒè§‚ç‚¹ï¼ˆå¦‚æœ‰ï¼‰ã€‚

åªè¿”å›æœ€ç»ˆç­”æ¡ˆã€‚"""

        response = self.llm_client.chat.completions.create(
            model=self.llm_name,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=600
        )
        
        return response.choices[0].message.content.strip()
    
    def _get_embedding(self, text: str) -> np.ndarray:
        """è·å–æ–‡æœ¬ embedding"""
        response = self.embedding_client.embeddings.create(
            model=self.embedding_name,
            input=text
        )
        return np.array(response.data[0].embedding, dtype='float32')
    
    # ==================== æŒä¹…åŒ– ====================
    
    def save(self, name: str = "default"):
        """ä¿å­˜çŸ¥è¯†åº“"""
        save_dir = self.storage_dir / name
        save_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"ğŸ’¾ ä¿å­˜çŸ¥è¯†åº“: {save_dir}")
        
        # ä¿å­˜æ–‡æ¡£
        with open(save_dir / "documents.json", 'w', encoding='utf-8') as f:
            json.dump(self.documents, f, ensure_ascii=False, indent=2)
        
        # ä¿å­˜ UUID æ˜ å°„
        with open(save_dir / "uuid_mappings.json", 'w', encoding='utf-8') as f:
            json.dump({
                'uuid_to_docid': self.uuid_to_docid,
                'docid_to_uuid': self.docid_to_uuid
            }, f, ensure_ascii=False, indent=2)
        
        # ä¿å­˜å›¾æ•°æ®
        with open(save_dir / "graph_data.pkl", 'wb') as f:
            pickle.dump({
                'text_chunks': self.text_chunks,
                'chunk_to_doc': self.chunk_to_doc,
                'entities': self.entities,
                'relationships': self.relationships,
                'claims': self.claims,
                'communities': self.communities,
                'community_summaries': self.community_summaries,
                'community_embeddings': self.community_embeddings,
            }, f)
        
        # ä¿å­˜å›¾
        with open(save_dir / "graph.gpickle", 'wb') as f:
            pickle.dump(self.graph, f)
        
        # ä¿å­˜ FAISS
        if self.community_summary_index:
            faiss.write_index(self.community_summary_index, 
                            str(save_dir / "faiss_index.bin"))
        
        logger.info(f"  âœ… ä¿å­˜å®Œæˆ: {len(self.documents)} æ–‡æ¡£, {len(self.text_chunks)} chunks")
    
    def load(self, name: str = "default"):
        """åŠ è½½çŸ¥è¯†åº“"""
        load_dir = self.storage_dir / name
        
        if not load_dir.exists():
            raise FileNotFoundError(f"çŸ¥è¯†åº“ä¸å­˜åœ¨: {load_dir}")
        
        logger.info(f"ğŸ“‚ åŠ è½½çŸ¥è¯†åº“: {load_dir}")
        
        # åŠ è½½æ–‡æ¡£
        with open(load_dir / "documents.json", 'r', encoding='utf-8') as f:
            self.documents = json.load(f)
        
        # åŠ è½½ UUID æ˜ å°„
        uuid_path = load_dir / "uuid_mappings.json"
        if uuid_path.exists():
            with open(uuid_path, 'r', encoding='utf-8') as f:
                mappings = json.load(f)
                self.uuid_to_docid = mappings['uuid_to_docid']
                self.docid_to_uuid = mappings['docid_to_uuid']
        
        # åŠ è½½å›¾æ•°æ®
        with open(load_dir / "graph_data.pkl", 'rb') as f:
            data = pickle.load(f)
            self.text_chunks = data['text_chunks']
            self.chunk_to_doc = data['chunk_to_doc']
            self.entities = data['entities']
            self.relationships = data['relationships']
            self.claims = data['claims']
            self.communities = data['communities']
            self.community_summaries = data['community_summaries']
            self.community_embeddings = data['community_embeddings']
        
        # åŠ è½½å›¾
        with open(load_dir / "graph.gpickle", 'rb') as f:
            self.graph = pickle.load(f)
        
        # åŠ è½½ FAISS
        index_path = load_dir / "faiss_index.bin"
        if index_path.exists():
            self.community_summary_index = faiss.read_index(str(index_path))
        
        logger.info(f"  âœ… åŠ è½½å®Œæˆ: {len(self.documents)} æ–‡æ¡£, "
                   f"{len(self.text_chunks)} chunks, {len(self.entities)} å®ä½“")